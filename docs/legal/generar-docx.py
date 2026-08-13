#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convierte los contratos en Markdown a .docx usando el membrete de Devifly.

Uso:
    python docs/legal/generar-docx.py docs/legal/servicios/*.md
    python docs/legal/generar-docx.py docs/legal/servicios/01-contrato-invitaciones-digitales.md
    python docs/legal/generar-docx.py --con-notas docs/legal/*.md
    python docs/legal/generar-docx.py --salida entregables docs/legal/servicios/*.md

Opciones:
    --con-notas   Incluye las cajas "> Cómo usar esta plantilla". Por omisión se
                  excluyen, porque el .docx es el documento que ve el cliente.
    --membrete    Ruta al .docx con el membrete. Por omisión:
                  ~/Downloads/membrete.docx
    --salida      Carpeta de salida. Por omisión: <carpeta del .md>/docx

Requiere: pip install python-docx
"""

import argparse
import glob
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("Falta python-docx. Instálalo con:  pip install python-docx")


# --- Paleta tomada del membrete -------------------------------------------

AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL_CLARO = RGBColor(0x2E, 0x74, 0xB5)
GRIS = RGBColor(0x77, 0x77, 0x77)
BORDE = "BFBFBF"
FONDO_ENCABEZADO = "EAEFF5"

FUENTE = "Calibri"
MEMBRETE_POR_OMISION = os.path.expanduser("~/Downloads/membrete.docx")


# --- Utilidades de OOXML ---------------------------------------------------

def _borde(elemento, color=BORDE, grosor="4"):
    elemento.set(qn("w:val"), "single")
    elemento.set(qn("w:sz"), grosor)
    elemento.set(qn("w:space"), "0")
    elemento.set(qn("w:color"), color)


def bordes_de_tabla(tabla):
    tblPr = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % lado)
        _borde(el)
        bordes.append(el)
    tblPr.append(bordes)


def sombrear_celda(celda, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(shd)


def linea_horizontal(parrafo):
    pPr = parrafo._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    _borde(bottom, color="D9D9D9", grosor="6")
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- Formato inline (**negritas**, *cursivas*, `código`) --------------------

TOKENS = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.S)
ENLACE = re.compile(r"\[([^\]]+)\]\((?:[^)]+)\)")


def escribir_inline(parrafo, texto, negrita_base=False, tam=None, color=None):
    texto = ENLACE.sub(r"\1", texto)
    texto = texto.replace("&nbsp;", " ")
    for parte in TOKENS.split(texto):
        if not parte:
            continue
        negrita, cursiva, mono = negrita_base, False, False
        if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
            parte, negrita = parte[2:-2], True
        elif parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            parte, cursiva = parte[1:-1], True
        elif parte.startswith("`") and parte.endswith("`") and len(parte) > 2:
            parte, mono = parte[1:-1], True
        run = parrafo.add_run(parte)
        run.bold = negrita
        run.italic = cursiva
        run.font.name = "Consolas" if mono else FUENTE
        if tam:
            run.font.size = tam
        if color:
            run.font.color.rgb = color


# --- Construcción del documento -------------------------------------------

def preparar_documento(ruta_membrete):
    doc = Document(ruta_membrete)

    cuerpo = doc.element.body
    for hijo in list(cuerpo):
        if not hijo.tag.endswith("}sectPr"):
            cuerpo.remove(hijo)

    seccion = doc.sections[0]
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.0)
    seccion.left_margin = Cm(2.5)
    seccion.right_margin = Cm(2.5)
    seccion.header_distance = Cm(1.25)
    seccion.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for nombre, tam, color, antes, despues in (
        ("Heading 1", 17, AZUL, 0, 10),
        ("Heading 2", 13, AZUL, 14, 6),
        ("Heading 3", 11.5, AZUL_CLARO, 12, 4),
        ("Heading 4", 11, AZUL_CLARO, 10, 4),
    ):
        estilo = doc.styles[nombre]
        estilo.font.name = FUENTE
        estilo.font.size = Pt(tam)
        estilo.font.bold = True
        estilo.font.color.rgb = color
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(despues)
        estilo.paragraph_format.keep_with_next = True

    return doc


def agregar_tabla(doc, filas):
    columnas = max(len(f) for f in filas)
    tabla = doc.add_table(rows=len(filas), cols=columnas)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = True
    bordes_de_tabla(tabla)

    for i, fila in enumerate(filas):
        for j in range(columnas):
            celda = tabla.cell(i, j)
            celda.paragraphs[0].text = ""
            contenido = fila[j] if j < len(fila) else ""
            # <br> dentro de una celda = renglones en blanco (bloques de firma)
            trozos = re.split(r"<br\s*/?>", contenido)
            for k, trozo in enumerate(trozos):
                p = celda.paragraphs[0] if k == 0 else celda.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                escribir_inline(p, trozo.strip(), negrita_base=(i == 0), tam=Pt(9.5))
            if i == 0:
                sombrear_celda(celda, FONDO_ENCABEZADO)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tabla


def convertir(ruta_md, doc, con_notas):
    with open(ruta_md, encoding="utf-8") as f:
        lineas = f.read().splitlines()

    i, n = 0, len(lineas)
    while i < n:
        linea = lineas[i]
        crudo = linea.rstrip()
        texto = crudo.strip()

        # Cajas de instrucciones internas
        if texto.startswith(">"):
            bloque = []
            while i < n and lineas[i].lstrip().startswith(">"):
                bloque.append(lineas[i].lstrip()[1:].strip())
                i += 1
            if con_notas:
                for parte in bloque:
                    if not parte:
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.6)
                    escribir_inline(p, parte, tam=Pt(9.5), color=GRIS)
            continue

        if not texto:
            i += 1
            continue

        if texto in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            linea_horizontal(p)
            i += 1
            continue

        if re.fullmatch(r"<br\s*/?>", texto):
            doc.add_paragraph()
            i += 1
            continue

        # Tablas GFM
        if texto.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lineas[i + 1].strip()):
            filas = []
            while i < n and lineas[i].strip().startswith("|"):
                fila = lineas[i].strip()
                if not re.match(r"^\|[\s:|-]+\|$", fila):
                    filas.append([c.strip() for c in fila.strip("|").split("|")])
                i += 1
            agregar_tabla(doc, filas)
            continue

        # Encabezados
        m = re.match(r"^(#{1,4})\s+(.*)$", texto)
        if m:
            nivel, contenido = len(m.group(1)), m.group(2)
            p = doc.add_paragraph(style="Heading %d" % nivel)
            if nivel == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            escribir_inline(p, contenido, negrita_base=True)
            i += 1
            continue

        # Casillas de verificación
        m = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)$", texto)
        if m:
            marca = "☒" if m.group(1).lower() == "x" else "☐"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(marca + "  ").font.name = "Segoe UI Symbol"
            escribir_inline(p, m.group(2))
            i += 1
            continue

        # Viñetas
        m = re.match(r"^[-*]\s+(.*)$", texto)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("•  ")
            escribir_inline(p, m.group(1))
            i += 1
            continue

        # Listas numeradas
        m = re.match(r"^(\d+)\.\s+(.*)$", texto)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("%s.  " % m.group(1))
            escribir_inline(p, m.group(2))
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escribir_inline(p, texto)
        i += 1

    return doc


def main():
    ap = argparse.ArgumentParser(description="Convierte contratos .md a .docx con el membrete de Devifly.")
    ap.add_argument("archivos", nargs="+", help="Archivos .md (acepta comodines)")
    ap.add_argument("--membrete", default=MEMBRETE_POR_OMISION, help="Ruta al .docx con el membrete")
    ap.add_argument("--salida", default=None, help="Carpeta de salida (por omisión: <carpeta del .md>/docx)")
    ap.add_argument("--con-notas", action="store_true", help="Conserva las cajas de instrucciones internas")
    args = ap.parse_args()

    if not os.path.isfile(args.membrete):
        sys.exit("No encontré el membrete en: %s" % args.membrete)

    rutas = []
    for patron in args.archivos:
        encontrados = glob.glob(patron)
        rutas.extend(encontrados if encontrados else [patron])
    rutas = [r for r in sorted(set(rutas)) if r.lower().endswith(".md")]

    if not rutas:
        sys.exit("No hay archivos .md que convertir.")

    for ruta in rutas:
        if not os.path.isfile(ruta):
            print("  ! no existe: %s" % ruta)
            continue
        carpeta = args.salida or os.path.join(os.path.dirname(os.path.abspath(ruta)), "docx")
        os.makedirs(carpeta, exist_ok=True)
        destino = os.path.join(carpeta, os.path.splitext(os.path.basename(ruta))[0] + ".docx")

        doc = preparar_documento(args.membrete)
        convertir(ruta, doc, args.con_notas)
        doc.save(destino)
        print("  OK  %s" % destino)


if __name__ == "__main__":
    main()
